from io import BytesIO
from pathlib import Path

from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.schemas.domain import CoverLetterPayload, ResumePayload


router = APIRouter()


def register_montserrat() -> str:
    font_path = Path("/Users/methode/Library/Fonts/Montserrat-Bold.otf")
    if "MontserratBold" in pdfmetrics.getRegisteredFontNames():
        return "MontserratBold"
    if font_path.exists():
        try:
            pdfmetrics.registerFont(TTFont("MontserratBold", str(font_path)))
            return "MontserratBold"
        except Exception:
            return "Helvetica-Bold"
    return "Helvetica-Bold"


def set_doc_font(run, font_name: str, size: float, bold: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)


def set_paragraph_spacing(paragraph, before: float = 0, after: float = 0, line: float = 1.15):
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    paragraph_format.line_spacing = line


def add_bottom_border(paragraph):
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def doc_styles():
    styles = getSampleStyleSheet()
    montserrat_font = register_montserrat()
    styles.add(
        ParagraphStyle(
            name="ResumeName",
            fontName=montserrat_font,
            fontSize=22,
            leading=24,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeContact",
            fontName=montserrat_font,
            fontSize=9.8,
            leading=13,
            alignment=TA_CENTER,
            textColor=colors.black,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SectionTitle",
            fontName=montserrat_font,
            fontSize=10.5,
            leading=12,
            alignment=TA_LEFT,
            textColor=colors.black,
            spaceAfter=2,
            spaceBefore=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Body",
            fontName=montserrat_font,
            fontSize=9.8,
            leading=14,
            alignment=TA_JUSTIFY,
            textColor=colors.black,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ItemHeading",
            fontName=montserrat_font,
            fontSize=10,
            leading=13,
            textColor=colors.black,
        )
    )
    styles.add(
        ParagraphStyle(
            name="Meta",
            fontName=montserrat_font,
            fontSize=9.6,
            leading=12.5,
            textColor=colors.black,
        )
    )
    styles.add(
        ParagraphStyle(
            name="SkillGroupTitle",
            fontName=montserrat_font,
            fontSize=9.5,
            leading=12,
            textColor=colors.black,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            name="ResumeBullet",
            fontName=montserrat_font,
            fontSize=9.8,
            leading=13.5,
            leftIndent=16,
            firstLineIndent=-10,
            textColor=colors.black,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterContact",
            fontName=montserrat_font,
            fontSize=10.2,
            leading=15,
            textColor=colors.black,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterSubject",
            fontName=montserrat_font,
            fontSize=11,
            leading=14,
            textColor=colors.black,
            spaceBefore=8,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="LetterBody",
            fontName=montserrat_font,
            fontSize=10.1,
            leading=18,
            alignment=TA_JUSTIFY,
            textColor=colors.black,
        )
    )
    return styles


def has_value(value) -> bool:
    return bool(str(value or "").strip())


def compact_entries(items: list[dict]) -> list[dict]:
    cleaned_items = []
    for item in items or []:
        if any(has_value(value) for value in item.values()):
            cleaned_items.append(item)
    return cleaned_items


def section_rule(title: str, styles) -> list:
    return [
        Spacer(1, 8),
        Paragraph(title.upper(), styles["SectionTitle"]),
        HRFlowable(width="100%", thickness=1, color=colors.black, spaceBefore=1, spaceAfter=6),
    ]


def bullet_points(raw_text: str, styles) -> list:
    points = []
    source = str(raw_text or "").replace("\r", "\n")
    chunks = []
    for line in source.split("\n"):
        for item in line.split("•"):
            chunks.extend(part for part in item.split(";"))
    for item in chunks:
        cleaned = item.strip()
        if cleaned:
            points.append(Paragraph(f"• {cleaned}", styles["ResumeBullet"]))
    return points


@router.post("/resume")
def build_resume(payload: ResumePayload) -> StreamingResponse:
    buffer = BytesIO()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    education_items = compact_entries(payload.education)
    experience_items = compact_entries(payload.experience)
    certification_items = compact_entries(payload.certifications)
    award_items = compact_entries(payload.awards)
    project_items = compact_entries(payload.projects)
    skill_group_items = compact_entries(payload.skill_groups)
    font_name = "Montserrat"

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=2, line=1.0)
    run = paragraph.add_run(payload.full_name)
    set_doc_font(run, font_name, 20, bold=True)

    for line in [
        " | ".join(part for part in [payload.location, payload.phone] if part),
        " | ".join(part for part in [
            payload.email,
            f"LinkedIn: {payload.linkedin_url}" if payload.linkedin_url else "",
            f"GitHub: {payload.github_url}" if payload.github_url else "",
        ] if part),
    ]:
        if line:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(paragraph, after=1, line=1.0)
            run = paragraph.add_run(line)
            set_doc_font(run, font_name, 10.4)

    def add_section_heading(title: str):
        paragraph = document.add_paragraph()
        set_paragraph_spacing(paragraph, before=7, after=3, line=1.0)
        run = paragraph.add_run(title.upper())
        set_doc_font(run, font_name, 10.8, bold=True)
        add_bottom_border(paragraph)

    add_section_heading("Professional Profile")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(paragraph, after=4, line=1.18)
    run = paragraph.add_run(payload.summary)
    set_doc_font(run, font_name, 10.2)

    if education_items:
        add_section_heading("Educational Background")
        for item in education_items:
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=1, line=1.0)
            run = paragraph.add_run(item.get("award", "Program"))
            set_doc_font(run, font_name, 10.8, bold=True)

            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=1, line=1.0)
            run = paragraph.add_run(" | ".join(part for part in [item.get("school", ""), item.get("location", ""), item.get("period", "")] if part))
            set_doc_font(run, font_name, 10.2)

            extra = " | ".join(part for part in [item.get("honors", ""), item.get("gpa", "")] if part)
            if extra:
                paragraph = document.add_paragraph()
                set_paragraph_spacing(paragraph, after=3, line=1.0)
                run = paragraph.add_run(extra)
                set_doc_font(run, font_name, 10.2)

    if experience_items:
        add_section_heading("Work Experience")
        for item in experience_items:
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=1, line=1.0)
            run = paragraph.add_run(", ".join(part for part in [item.get("period", ""), item.get("role", ""), item.get("company", ""), item.get("location", "")] if part))
            set_doc_font(run, font_name, 10.8, bold=True)

            for bullet in [part.strip() for part in str(item.get("highlights", "")).replace("\r", "\n").replace("•", "\n").split("\n") if part.strip()]:
                for chunk in [part.strip() for part in bullet.split(";") if part.strip()]:
                    bullet_paragraph = document.add_paragraph(style=None)
                    bullet_paragraph.paragraph_format.left_indent = Inches(0.18)
                    set_paragraph_spacing(bullet_paragraph, after=0.5, line=1.0)
                    run = bullet_paragraph.add_run(f"• {chunk}")
                    set_doc_font(run, font_name, 10.2)

            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=2)

    cert_and_awards = []
    for item in award_items:
        cert_and_awards.append(f"{item.get('name', 'Award')} | {item.get('issuer', '')} | {item.get('year', '')}")
    for item in certification_items:
        cert_and_awards.append(f"{item.get('name', 'Certification')} | {item.get('issuer', '')} | {item.get('year', '')}")
    if cert_and_awards:
        add_section_heading("Certifications & Awards")
        for entry in cert_and_awards:
            paragraph = document.add_paragraph()
            bullet = paragraph.add_run("• ")
            set_doc_font(bullet, font_name, 10.2)
            run = paragraph.add_run(entry)
            set_doc_font(run, font_name, 10.2, bold=False)
            set_paragraph_spacing(paragraph, after=0.5, line=1.0)

    if skill_group_items:
        add_section_heading("Skills & Abilities")
        skill_groups = skill_group_items[:6]
        rows_needed = (len(skill_groups) + 2) // 3
        table = document.add_table(rows=rows_needed, cols=3)
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(2.18)

        for index, group in enumerate(skill_groups):
            cell = table.rows[index // 3].cells[index % 3]
            title = cell.paragraphs[0]
            set_paragraph_spacing(title, after=0.5, line=1.0)
            run = title.add_run(f"{group.get('name', 'Skill Group')}:")
            set_doc_font(run, font_name, 10.5, bold=True)

            for skill in [part.strip() for part in str(group.get("skills", "")).split(",") if part.strip()]:
                paragraph = cell.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.1)
                set_paragraph_spacing(paragraph, after=0.5, line=1.0)
                run = paragraph.add_run(f"• {skill}")
                set_doc_font(run, font_name, 10.2)

    if project_items:
        add_section_heading("Projects")
        for item in project_items:
            heading = " | ".join(part for part in [item.get("name", "Project"), item.get("impact", ""), item.get("link", "")] if part)
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=1, line=1.0)
            run = paragraph.add_run(heading)
            set_doc_font(run, font_name, 10.8, bold=True)

            if item.get("description"):
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_paragraph_spacing(paragraph, after=3, line=1.15)
                run = paragraph.add_run(item.get("description", ""))
                set_doc_font(run, font_name, 10.2)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=6, after=1, line=1.0)
    add_bottom_border(paragraph)
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=1, line=1.0)
    run = paragraph.add_run("References Available Upon Request")
    set_doc_font(run, font_name, 10.8, bold=True)
    add_bottom_border(paragraph)

    document.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="alx-resume.docx"'},
    )


@router.post("/cover-letter")
def build_cover_letter(payload: CoverLetterPayload) -> StreamingResponse:
    buffer = BytesIO()
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    font_name = "Montserrat"

    for line in [
        payload.full_name,
        payload.location,
        payload.phone,
        payload.email,
        payload.linkedin_url,
    ]:
        if line:
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=0.5, line=1.0)
            run = paragraph.add_run(str(line))
            set_doc_font(run, font_name, 10.8, bold=False)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=6)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=1, line=1.0)
    run = paragraph.add_run(f"{payload.date},")
    set_doc_font(run, font_name, 10.8)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=6)

    for line in [
        payload.hiring_manager,
        payload.recipient_title,
        payload.company,
        payload.recipient_address,
        payload.recipient_location,
    ]:
        if line:
            paragraph = document.add_paragraph()
            set_paragraph_spacing(paragraph, after=0.5, line=1.0)
            run = paragraph.add_run(str(line))
            set_doc_font(run, font_name, 10.8)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=8, after=8, line=1.0)
    run = paragraph.add_run(f"RE: {payload.role_title}".upper())
    set_doc_font(run, font_name, 12.8, bold=True)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=8, line=1.0)
    run = paragraph.add_run("Dear Hiring Team,")
    set_doc_font(run, font_name, 10.8)

    for block in [payload.intro, *payload.body, payload.closing]:
        if block:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(paragraph, after=7, line=1.18)
            run = paragraph.add_run(str(block))
            set_doc_font(run, font_name, 10.8)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, before=2, after=8, line=1.0)
    run = paragraph.add_run("Sincerely,")
    set_doc_font(run, font_name, 10.8)

    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=0)
    run = paragraph.add_run(payload.full_name)
    set_doc_font(run, font_name, 12.2, bold=True)

    document.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="alx-cover-letter.docx"'},
    )
